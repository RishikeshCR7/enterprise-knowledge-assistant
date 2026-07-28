import os
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
import docx
import openpyxl

DATA_DIR = "/Users/pooja/Documents/GitHub/enterprise-knowledge-assistant/data"

DEPARTMENTS = ["HR", "Engineering", "Finance", "Legal", "Sales"]


def create_pdf(file_path: str, title: str, paragraphs: list[str]):
    os.makedirs(os.path.dirname(file_path), exist_ok=True)
    doc = SimpleDocTemplate(file_path, pagesize=letter)
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        'DocTitle',
        parent=styles['Heading1'],
        fontSize=18,
        spaceAfter=14
    )
    body_style = ParagraphStyle(
        'DocBody',
        parent=styles['Normal'],
        fontSize=11,
        leading=15,
        spaceAfter=10
    )
    story = [Paragraph(title, title_style), Spacer(1, 10)]
    for p in paragraphs:
        story.append(Paragraph(p, body_style))
        story.append(Spacer(1, 8))
    doc.build(story)
    print(f"Created PDF: {file_path}")


def create_docx(file_path: str, title: str, paragraphs: list[str]):
    os.makedirs(os.path.dirname(file_path), exist_ok=True)
    doc = docx.Document()
    doc.add_heading(title, 0)
    for p in paragraphs:
        doc.add_paragraph(p)
    doc.save(file_path)
    print(f"Created DOCX: {file_path}")


def create_xlsx(file_path: str, title: str, headers: list[str], rows: list[list]):
    os.makedirs(os.path.dirname(file_path), exist_ok=True)
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = title[:30]
    ws.append(headers)
    for row in rows:
        ws.append(row)
    wb.save(file_path)
    print(f"Created XLSX: {file_path}")


def generate_all():
    # HR Documents
    create_pdf(
        os.path.join(DATA_DIR, "HR", "LeavePolicy.pdf"),
        "Enterprise Human Resources Leave Policy 2026",
        [
            "1. Overview: The Enterprise Human Resources Leave Policy governs paid time off, sick leave, parental leave, and emergency absences for all full-time employees.",
            "2. Annual Leave Entitlement: Employees are entitled to 20 business days of paid annual leave per calendar year, accrued monthly at a rate of 1.66 days.",
            "3. Sick Leave: Full-time employees receive 10 days of paid sick leave annually. Absences exceeding 3 consecutive days require a physician note.",
            "4. Parental Leave: Eligible employees with at least 12 months of service receive 12 weeks of fully paid parental leave following the birth or adoption of a child.",
            "5. Approval Workflow: All leave requests must be submitted through the Employee Portal at least 14 days in advance and approved by the department manager."
        ]
    )

    create_docx(
        os.path.join(DATA_DIR, "HR", "SalaryPolicy.docx"),
        "Compensation and Salary Structure Guidelines",
        [
            "Section 1: Compensation Philosophy. The company aims to provide competitive market-based compensation tied to performance and annual benchmark surveys.",
            "Section 2: Performance Reviews. Annual salary adjustments occur in Q1 following performance evaluations. Base salaries are benchmarked against 75th percentile industry averages.",
            "Section 3: Bonus Eligibility. Discretionary performance bonuses are distributed semi-annually based on company earnings and individual KPI achievement.",
            "Section 4: Confidentiality. Salary details are confidential between the employee, HR, and direct executive management."
        ]
    )

    create_pdf(
        os.path.join(DATA_DIR, "HR", "HiringProcess.pdf"),
        "Standard Operating Procedure: Talent Acquisition and Hiring",
        [
            "1. Requisition: Hiring managers initiate job requisitions via HRIS. All postings require budget approval from the Finance department.",
            "2. Interview Stages: The hiring process consists of: (a) Recruiter Screening, (b) Technical/Role Assessment, (c) Panel Interview, (d) Executive Approval.",
            "3. Background Verification: All conditional offers are contingent upon background checks, education verification, and reference validation.",
            "4. Onboarding: New hires complete orientation on Day 1, receiving security badges, IT assets, and department onboarding schedules."
        ]
    )

    # Engineering Documents
    create_pdf(
        os.path.join(DATA_DIR, "Engineering", "CodingStandards.pdf"),
        "Software Engineering Coding Standards & Architecture Guidelines",
        [
            "1. Python Standards: Follow PEP 8 guidelines. All code must pass Black formatting, Flake8 linting, and MyPy type checking before PR submission.",
            "2. Microservices Architecture: Services must communicate asynchronously via message queues or RESTful APIs with Pydantic contract validation.",
            "3. Security Practices: Never commit hardcoded credentials or API keys. Use environment variables and secrets manager vaults.",
            "4. Unit Testing: Maintain a minimum of 80% test coverage across core application logic. Integration tests must be automated in CI/CD."
        ]
    )

    create_docx(
        os.path.join(DATA_DIR, "Engineering", "API_Guidelines.docx"),
        "REST API Specification and Versioning Guidelines",
        [
            "API Conventions: All external facing HTTP endpoints must be versioned using URI path prefixes (e.g., /api/v1/resource).",
            "Request & Response Formatting: Use JSON payloads with standardized error responses containing status code, error code, and detailed message.",
            "Authentication: Protect endpoints using JWT tokens with Bearer authorization headers. Validate scope permissions per route.",
            "Rate Limiting: Endpoints enforce standard rate limits of 100 requests per minute per IP address."
        ]
    )

    create_pdf(
        os.path.join(DATA_DIR, "Engineering", "DockerGuide.pdf"),
        "Containerization and Deployment Guide",
        [
            "1. Base Images: Standardize on slim official Linux images (python:3.11-slim or node:20-alpine) to reduce security vulnerability surfaces.",
            "2. Multi-Stage Builds: Leverage multi-stage Dockerfiles to keep production container images under 200MB.",
            "3. Security Scanning: Docker images must be scanned for vulnerabilities via Trivy prior to registry push.",
            "4. Kubernetes Deployment: Deploy services using Helm charts configured with resource limits and readiness probes."
        ]
    )

    # Finance Documents
    create_pdf(
        os.path.join(DATA_DIR, "Finance", "ExpensePolicy.pdf"),
        "Corporate Travel and Business Expense Reimbursement Policy",
        [
            "1. Scope: Applies to all business travel, client entertainment, and operational purchases incurred on behalf of the company.",
            "2. Travel Expenses: Flight bookings must be economy class for domestic flights under 6 hours. Lodging expenses must not exceed $250 per night.",
            "3. Per Diem Meals: Daily meal allowance is capped at $75 per day. Itemized receipts are required for all individual expenses exceeding $25.",
            "4. Submission Timeline: Expense reports must be submitted within 30 days of travel via the Finance Expense Portal."
        ]
    )

    create_xlsx(
        os.path.join(DATA_DIR, "Finance", "QuarterlyBudget.xlsx"),
        "Quarterly Allocation",
        ["Department", "Q1 Budget", "Q2 Budget", "Q3 Budget", "Q4 Budget", "Approved By"],
        [
            ["HR", 150000, 160000, 155000, 170000, "VP Finance"],
            ["Engineering", 850000, 900000, 950000, 1000000, "CFO"],
            ["Finance", 200000, 210000, 205000, 220000, "CFO"],
            ["Legal", 300000, 310000, 305000, 320000, "VP Finance"],
            ["Sales", 500000, 550000, 600000, 650000, "CFO"]
        ]
    )

    # Legal Documents
    create_pdf(
        os.path.join(DATA_DIR, "Legal", "VendorAgreement.pdf"),
        "Master Service Agreement and Vendor Legal Compliance",
        [
            "1. Contract Term: Master Service Agreements carry a standard 12-month term automatically renewing unless written notice is given 60 days prior.",
            "2. Data Protection & GDPR: Vendors processing customer data must comply with GDPR, CCPA, and sign standard Data Processing Addendums (DPA).",
            "3. Indemnification: Vendors agree to defend and hold harmless the company against third-party IP infringement claims.",
            "4. Governing Law: This agreement shall be governed by and construed under the laws of the State of Delaware."
        ]
    )

    create_pdf(
        os.path.join(DATA_DIR, "Legal", "Compliance.pdf"),
        "Corporate Governance and Security Compliance Manual",
        [
            "1. SOC 2 Type II Compliance: The organization strictly adheres to SOC 2 Security, Confidentiality, and Availability trust service criteria.",
            "2. Access Controls: Role-based access control (RBAC) must be audited quarterly to enforce least-privilege principles across all systems.",
            "3. Incident Response: Security incidents must be reported to security@company.com within 1 hour of detection.",
            "4. Data Retention: Records must be retained for 7 years before secure destruction according to legal retention schedules."
        ]
    )

    # Sales Documents
    create_pdf(
        os.path.join(DATA_DIR, "Sales", "SalesStrategy.pdf"),
        "Enterprise Sales Strategy and Market Expansion 2026",
        [
            "1. Target Audience: Focus sales acquisition on Fortune 500 enterprises needing secure private RAG knowledge assistants.",
            "2. Value Proposition: Highlight zero-data-retention, SOC 2 compliance, local vector search, and seamless RBAC integration.",
            "3. Sales Cycle: Standard enterprise deal cycle spans 90 to 120 days from initial demo to contract signature.",
            "4. Pricing Tiers: Base platform license starts at $50,000/year with tier discounts for additional user seats."
        ]
    )

    create_docx(
        os.path.join(DATA_DIR, "Sales", "ClientOnboarding.docx"),
        "Client Onboarding and Implementation Guide",
        [
            "Phase 1: Kickoff. Conduct kickoff call with client executive sponsors and technical leads within 5 business days of contract execution.",
            "Phase 2: Data Ingestion. Configure data connectors to index client documents, setting up security boundaries and role mappings.",
            "Phase 3: User Training. Deliver admin and end-user training sessions, providing custom user documentation.",
            "Phase 4: Go-Live. Conduct final verification and launch solution into production with dedicated account management support."
        ]
    )


if __name__ == "__main__":
    generate_all()
    print("Sample enterprise dataset generated successfully!")
